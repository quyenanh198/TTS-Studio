"""Parser tests with synthetic fixtures generated on the fly."""

from pathlib import Path

import pytest

from app.services import parsers

LOREM = "Đây là một câu văn mẫu dùng để kiểm thử bộ tách chương của ứng dụng. " * 6


def _book_text(n=3):
    parts = ["Lời tựa\n" + LOREM]
    for i in range(1, n + 1):
        parts.append(f"Chương {i}: Tiêu đề {i}\n{LOREM}\n\n{LOREM}")
    return "\n\n".join(parts)


def test_split_by_headings_vietnamese():
    book = parsers.parse_plain(_book_text(3), "Sách")
    titles = [c.title for c in book.chapters]
    assert titles[0] == "Mở đầu"
    assert titles[1:] == ["Chương 1: Tiêu đề 1", "Chương 2: Tiêu đề 2", "Chương 3: Tiêu đề 3"]
    assert all(len(c.text) > 100 for c in book.chapters)


def test_single_chapter_when_no_headings():
    book = parsers.parse_plain(LOREM, "Đơn")
    assert len(book.chapters) == 1 and book.chapters[0].title == "Nội dung"


def test_heading_patterns():
    assert parsers.is_heading("CHAPTER 12 - The End")
    assert parsers.is_heading("第三章 风起")
    assert parsers.is_heading("Hồi 5")
    assert not parsers.is_heading("Chương trình này rất hay và dài " * 5)
    assert not parsers.is_heading("Anh ấy nói: chương này khó.")


def test_txt_file(tmp_path: Path):
    p = tmp_path / "s.txt"
    p.write_text(_book_text(2), encoding="utf-8")
    book = parsers.parse_file(p)
    assert book.format == "txt" and len(book.chapters) == 3


def test_srt_file(tmp_path: Path):
    p = tmp_path / "s.srt"
    p.write_text("1\n00:00:00,000 --> 00:00:01,000\nHello\n\n2\n00:00:01,500 --> 00:00:03,000\nWorld\n", encoding="utf-8")
    book = parsers.parse_file(p)
    ch = book.chapters[0]
    assert ch.cues and len(ch.cues) == 2 and ch.cues[1]["start"] == 1.5
    assert ch.text == "Hello\nWorld"


def test_docx_file(tmp_path: Path):
    import docx

    d = docx.Document()
    d.add_heading("Chương 1 Bắt đầu", level=1)
    d.add_paragraph(LOREM)
    d.add_heading("Chương 2 Tiếp theo", level=1)
    d.add_paragraph(LOREM)
    p = tmp_path / "b.docx"
    d.save(str(p))
    book = parsers.parse_file(p)
    assert [c.title for c in book.chapters] == ["Chương 1 Bắt đầu", "Chương 2 Tiếp theo"]


def test_epub_file(tmp_path: Path):
    from ebooklib import epub

    bk = epub.EpubBook()
    bk.set_identifier("id1")
    bk.set_title("Sách EPUB")
    bk.set_language("vi")
    items = []
    for i in (1, 2):
        c = epub.EpubHtml(title=f"Chương {i}", file_name=f"c{i}.xhtml", lang="vi")
        c.content = f"<h1>Chương {i}</h1><p>{LOREM}</p><p>{LOREM}</p>"
        bk.add_item(c)
        items.append(c)
    bk.toc = tuple(items)
    bk.add_item(epub.EpubNcx())
    bk.add_item(epub.EpubNav())
    bk.spine = ["nav", *items]
    p = tmp_path / "b.epub"
    epub.write_epub(str(p), bk)
    book = parsers.parse_file(p)
    assert book.title == "Sách EPUB"
    assert [c.title for c in book.chapters] == ["Chương 1", "Chương 2"]
    assert LOREM.strip()[:30] in book.chapters[0].text


def test_pdf_file(tmp_path: Path):
    import fitz

    doc = fitz.open()
    for i in (1, 2):
        page = doc.new_page()
        page.insert_text((72, 72), f"Chapter {i} Title", fontsize=16)
        y = 110
        for _ in range(8):
            page.insert_text((72, y), "This is a sample sentence for the parser test.", fontsize=11)
            y += 16
    p = tmp_path / "b.pdf"
    doc.save(str(p))
    book = parsers.parse_file(p)
    assert book.format == "pdf"
    assert len(book.chapters) == 2
    assert book.chapters[0].title.startswith("Chapter 1")


def test_unsupported(tmp_path: Path):
    p = tmp_path / "x.xyz"
    p.write_text("a")
    with pytest.raises(ValueError):
        parsers.parse_file(p)
